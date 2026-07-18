from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
PNG_PATH = OUT / "visioals-personalized-response-flowchart.png"
DOCX_PATH = OUT / "visioals-personalized-response-overview.docx"

FONT = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"

NAVY = "17365D"
BLUE = "2E75B6"
TEAL = "168C8C"
PURPLE = "6B5CA5"
GOLD = "C28B2C"
INK = "1F2937"
MUTED = "5B6573"
LOCAL_FILL = "E8F3FA"
CLOUD_FILL = "F1ECFA"
DATA_FILL = "E8F6F2"
USER_FILL = "FFF4DD"
LANE_FILL = "F7F9FC"
WHITE = "FFFFFF"
LINE = "CBD5E1"


def rgb(hex_color):
    return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))


def font(size, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT, size)


def rounded(draw, box, fill, outline=LINE, radius=22, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=rgb(fill), outline=rgb(outline), width=width)


def centered_text(draw, box, title, detail="", title_size=29, detail_size=20,
                  color=INK, title_color=None):
    x1, y1, x2, y2 = box
    tf = font(title_size, True)
    df = font(detail_size, False)
    lines = [title]
    detail_lines = detail.split("\n") if detail else []
    title_bbox = draw.textbbox((0, 0), title, font=tf)
    title_h = title_bbox[3] - title_bbox[1]
    detail_h = sum((draw.textbbox((0, 0), s, font=df)[3] - draw.textbbox((0, 0), s, font=df)[1]) for s in detail_lines)
    gap = 12 if detail_lines else 0
    line_gap = 6 * max(len(detail_lines) - 1, 0)
    total_h = title_h + gap + detail_h + line_gap
    y = y1 + (y2 - y1 - total_h) / 2
    tw = title_bbox[2] - title_bbox[0]
    draw.text(((x1 + x2 - tw) / 2, y), title, font=tf, fill=rgb(title_color or color))
    y += title_h + gap
    for s in detail_lines:
        bb = draw.textbbox((0, 0), s, font=df)
        w, h = bb[2] - bb[0], bb[3] - bb[1]
        draw.text(((x1 + x2 - w) / 2, y), s, font=df, fill=rgb(color))
        y += h + 6


def arrow(draw, start, end, color=BLUE, width=7):
    draw.line([start, end], fill=rgb(color), width=width)
    x2, y2 = end
    x1, y1 = start
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - direction * 20, y2 - 14), (x2 - direction * 20, y2 + 14)]
    else:
        direction = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 14, y2 - direction * 20), (x2 + 14, y2 - direction * 20)]
    draw.polygon(pts, fill=rgb(color))


def create_diagram():
    canvas = Image.new("RGB", (2200, 1120), rgb(WHITE))
    d = ImageDraw.Draw(canvas)

    lane_x1, lane_x2 = 42, 2158
    lanes = [
        (36, 338, "1  BUILD THE PATIENT PROFILE", "Set up or refresh when sources change"),
        (374, 820, "2  GENERATE A RESPONSE", "Runs for every caregiver question"),
        (856, 1084, "3  ADAPT OVER TIME", "Selection and rejection history shapes future options"),
    ]
    for y1, y2, label, sub in lanes:
        rounded(d, (lane_x1, y1, lane_x2, y2), LANE_FILL, outline="DFE6EF", radius=24, width=2)
        d.text((70, y1 + 18), label, font=font(25, True), fill=rgb(NAVY))
        d.text((70, y1 + 52), sub, font=font(18), fill=rgb(MUTED))

    # Lane 1: profile setup
    boxes1 = [
        ((80, 132, 345, 292), "Patient sources", "Writing · audio · video\npasted text", DATA_FILL, TEAL),
        ((395, 132, 675, 292), "Local preparation", "Load text · transcribe media\nwith Whisper", LOCAL_FILL, BLUE),
        ((760, 132, 1065, 292), "Patient corpus", "Written + spoken\nlanguage samples", DATA_FILL, TEAL),
        ((1160, 104, 1520, 200), "Style profile", "spaCy · TF-IDF · heuristics · LLM", LOCAL_FILL, BLUE),
        ((1160, 218, 1520, 314), "Semantic index", "MiniLM embeddings · cosine similarity", LOCAL_FILL, BLUE),
        ((1645, 132, 2075, 292), "Personalization assets", "Style summary · corpus vectors\nvoice ID · preference profile", DATA_FILL, TEAL),
    ]
    for box, title, detail, fill, accent in boxes1:
        rounded(d, box, fill, outline=accent)
        centered_text(d, box, title, detail, title_size=27, detail_size=18, title_color=accent)
    arrow(d, (345, 212), (383, 212), BLUE)
    arrow(d, (675, 212), (748, 212), BLUE)
    arrow(d, (1065, 190), (1148, 154), BLUE)
    arrow(d, (1065, 234), (1148, 266), BLUE)
    arrow(d, (1520, 152), (1633, 188), TEAL)
    arrow(d, (1520, 266), (1633, 230), TEAL)

    # Lane 2: live pipeline
    boxes2 = [
        ((80, 490, 340, 690), "Caregiver asks", "Microphone audio", USER_FILL, GOLD),
        ((405, 490, 700, 690), "Question text", "Local Whisper\ntiny.en · CPU INT8", LOCAL_FILL, BLUE),
        ((770, 452, 1125, 728), "Assemble context", "Style summary\nRelevant corpus examples\nRecent conversation\nPreference rules\nRejected options", DATA_FILL, TEAL),
        ((1195, 490, 1510, 690), "Cloud LLM", "GPT-5.6 Luna\nvia OpenRouter", CLOUD_FILL, PURPLE),
        ((1580, 490, 1840, 690), "4 short options", "2–8 words each", USER_FILL, GOLD),
        ((1905, 490, 2150, 690), "Gaze selection", "Patient retains\nfinal control", USER_FILL, GOLD),
    ]
    for box, title, detail, fill, accent in boxes2:
        rounded(d, box, fill, outline=accent)
        centered_text(d, box, title, detail, title_size=27, detail_size=18, title_color=accent)
    for a, b, c in [
        ((340, 590), (393, 590), BLUE), ((700, 590), (758, 590), TEAL),
        ((1125, 590), (1183, 590), PURPLE), ((1510, 590), (1568, 590), PURPLE),
        ((1840, 590), (1893, 590), GOLD),
    ]:
        arrow(d, a, b, c)

    # Expansion and speech subflow
    rounded(d, (1195, 724, 1510, 790), CLOUD_FILL, outline=PURPLE, radius=18)
    centered_text(d, (1195, 724, 1510, 790), "Brief LLM expansion", "", 22, 17, title_color=PURPLE)
    rounded(d, (1580, 724, 2150, 790), LOCAL_FILL, outline=BLUE, radius=18)
    centered_text(d, (1580, 724, 2150, 790), "ElevenLabs cloned voice  →  local TTS fallback", "", 21, 17, title_color=BLUE)
    arrow(d, (2028, 690), (2028, 712), GOLD, width=6)
    d.line([(2028, 712), (1354, 712), (1354, 724)], fill=rgb(PURPLE), width=6)
    arrow(d, (1510, 757), (1568, 757), BLUE, width=6)

    # Lane 3: adaptation
    boxes3 = [
        ((160, 940, 620, 1036), "Log selections + rejections", "Stored locally as JSONL", DATA_FILL, TEAL),
        ((815, 940, 1330, 1036), "Every 20 interactions", "LLM extracts 3–8 preference rules", CLOUD_FILL, PURPLE),
        ((1530, 940, 2035, 1036), "Improve future candidates", "Rules return to the context package", DATA_FILL, TEAL),
    ]
    for box, title, detail, fill, accent in boxes3:
        rounded(d, box, fill, outline=accent, radius=18)
        centered_text(d, box, title, detail, title_size=23, detail_size=17, title_color=accent)
    arrow(d, (620, 988), (803, 988), PURPLE, width=6)
    arrow(d, (1330, 988), (1518, 988), TEAL, width=6)

    canvas.save(PNG_PATH, quality=95, dpi=(200, 200))


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def set_run(run, size, color=INK, bold=False, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def create_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(11)
    sec.page_height = Inches(8.5)
    sec.top_margin = Inches(0.42)
    sec.bottom_margin = Inches(0.40)
    sec.left_margin = Inches(0.50)
    sec.right_margin = Inches(0.50)
    sec.header_distance = Inches(0.25)
    sec.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(title.add_run("How VisioALS Personalizes a Response"), 23, NAVY, True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(5)
    set_run(subtitle.add_run(
        "High-level architecture: local language processing + semantic retrieval + cloud generation + patient-controlled selection"
    ), 10.5, MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(PNG_PATH), width=Inches(10.0), height=Inches(5.09))
    drawing = run._element.xpath(".//wp:docPr")
    if drawing:
        drawing[0].set("descr", "Flowchart of the VisioALS personalized-response pipeline")

    # Key takeaway band, true one-row table with fixed geometry.
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(10)
    cell = table.cell(0, 0)
    cell.width = Inches(10)
    set_cell_shading(cell, "EEF4FA")
    set_cell_margins(cell, top=90, bottom=90, start=150, end=150)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(0)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(cp.add_run("KEY IDEA  "), 10, BLUE, True)
    set_run(cp.add_run(
        "The patient does not train a new language model. VisioALS conditions pretrained models with a style profile, retrieved examples, and learned preference rules."
    ), 10, INK)

    legend = doc.add_paragraph()
    legend.paragraph_format.space_before = Pt(4)
    legend.paragraph_format.space_after = Pt(0)
    legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, color, desc in [
        ("LOCAL PROCESSING", BLUE, "runs on the device"),
        ("PATIENT DATA", TEAL, "stored personalization assets"),
        ("CLOUD AI", PURPLE, "LLM or cloned voice service"),
        ("PATIENT CONTROL", GOLD, "human question, choice, or output"),
    ]:
        set_run(legend.add_run(f"  ■ {label} "), 7.8, color, True)
        set_run(legend.add_run(f"{desc}   "), 7.8, MUTED)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    set_run(footer.add_run("VisioALS · Personalized Response Pipeline · July 2026"), 8, MUTED)

    # Fixed table width and indent in OOXML.
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "14400")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    doc.core_properties.title = "How VisioALS Personalizes a Response"
    doc.core_properties.subject = "High-level flowchart of the VisioALS personalized-response pipeline"
    doc.core_properties.author = "VisioALS"
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    create_diagram()
    create_docx()
    print(DOCX_PATH)
